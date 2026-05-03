"""
Generate PartnerAI Project Report as Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

#  Page margins 
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

#  Styles 
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    hs.font.name = 'Calibri'

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f"  {text}")
    else:
        p.add_run(text)

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(1)


# 
# TITLE PAGE
# 
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PartnerAI')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x6B, 0x5B, 0x95)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI-Powered Personal Productivity & Growth Companion')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run('Full Project Report\n')
run.font.size = Pt(14)
run = details.add_run('Technical Documentation, Architecture & RLHF Fine-Tuning')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Date: February 2026\n')
run.font.size = Pt(11)
run = info.add_run('Tech Stack: Python Flask  Phi-3 LLM  SQLite  APScheduler\n')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = info.add_run('Architecture: Local-First AI with RLHF Strategy Selection')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# 
# TABLE OF CONTENTS
# 
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Tech Stack',
    '3. System Architecture',
    '4. AI Fine-Tuning  RLHF System (Step by Step)',
    '5. Dynamic Persona System',
    '6. Automation System',
    '7. Reminder & Notification System',
    '8. Database Schema',
    '9. Frontend Pages',
    '10. Chat Command System',
    '11. How to Run the Project',
    '12. Key Design Decisions',
    '13. File Structure Reference',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# 
# 1. PROJECT OVERVIEW
# 
doc.add_heading('1. Project Overview', level=1)

doc.add_paragraph(
    'PartnerAI is an AI-powered personal productivity and growth companion that combines '
    'habit tracking, intelligent task management, team collaboration, and personalized coaching. '
    'The system uses a local AI model (Microsoft Phi-3 via llama.cpp) to provide privacy-focused '
    'mentorship through a comprehensive web dashboard.'
)

doc.add_heading('Key Features', level=2)
features = [
    ('Personalized AI Mentor', 'Natural conversation with context-aware coaching based on user goals'),
    ('Smart Task Management', 'AI-generated daily tasks, tracking, and automatic rollover'),
    ('Habit Tracking', 'Create habits, log completions, AI-powered insights and analytics'),
    ('Email Reminders', 'Scheduled reminders via Gmail SMTP at user-specified times'),
    ('In-App Notifications', 'Real-time reminder popups with polling system'),
    ('RLHF Feedback Loop', 'Thumbs up/down on AI responses adapts coaching style over time'),
    ('Focus Mode', 'Pomodoro timer with ambient music and wallpapers'),
    ('Knowledge Blocks', 'Notion-like workspace for notes, ideas, and learning'),
    ('Team Collaboration', 'Groups, shared tasks, team chat, invite codes'),
    ('Weekly Reports', 'AI-generated coaching reports with strengths/weaknesses'),
    ('Poll Suggestions', 'Smart clickable quick-reply buttons after every AI response'),
    ('Gamification', 'Rewards, streaks, and badges for consistency'),
]
for title, desc in features:
    add_bullet(desc, title)

doc.add_heading('Target Users', level=2)
add_bullet('Individuals seeking personal growth, skill development, or productivity improvement')
add_bullet('Students managing study goals and exam preparation')
add_bullet('Developers and professionals tracking learning progress')
add_bullet('Small teams working on collaborative projects')

doc.add_page_break()

# 
# 2. TECH STACK
# 
doc.add_heading('2. Tech Stack', level=1)

add_table(
    ['Layer', 'Technology', 'Details'],
    [
        ['Frontend', 'HTML/CSS/JS + Jinja2', 'Vanilla JS, no framework, dark theme design system'],
        ['Backend', 'Python Flask', 'Single-file server (app.py ~3300 lines), REST API'],
        ['Database', 'SQLite', 'partnerai.db  18+ tables, thread-safe connections'],
        ['AI Model', 'Microsoft Phi-3-mini-4k', 'Q4 quantized GGUF format, 3.8B parameters'],
        ['AI Server', 'llama.cpp (llama-server.exe)', 'Runs on http://127.0.0.1:8080, OpenAI-compatible API'],
        ['Scheduler', 'APScheduler', 'BackgroundScheduler with CronTrigger for 3 daily jobs'],
        ['Email', 'Gmail SMTP (smtplib)', 'TLS on port 587, app-specific password'],
        ['RLHF', 'Custom Python module', 'Epsilon-greedy strategy selection with feedback logging'],
        ['Parsing', 'Custom NLP (regex)', 'Natural language time parsing for reminders'],
    ]
)

doc.add_heading('Python Dependencies (requirements.txt)', level=2)
deps = ['flask  Web framework', 'requests  HTTP client for llama-server API',
        'apscheduler  Background job scheduling', 'google-generativeai  Fallback AI (optional)',
        'ollama  Alternative LLM interface (optional)']
for d in deps:
    add_bullet(d)

doc.add_page_break()

# 
# 3. SYSTEM ARCHITECTURE
# 
doc.add_heading('3. System Architecture', level=1)

doc.add_heading('3.1 High-Level Architecture', level=2)
doc.add_paragraph(
    'The system follows a layered architecture with clear separation between '
    'user interface, application logic, intelligence engines, and data persistence.'
)

add_code_block(
    'User (Browser)\n'
    '    \n'
    '    \n'
    'Flask Web Server (app.py, port 5000)\n'
    '    \n'
    '     Templates (HTML/CSS/JS)  Rendered pages\n'
    '     SQLite (memory.py)  User data, chat, tasks, habits\n'
    '     Local LLM (local_llm.py)  Phi-3 via llama-server:8080\n'
    '     RLHF (rlhf/)  Feedback  strategy selection\n'
    '     Scheduler (ai_task_scheduler.py)  Cron jobs\n'
    '     Email (smtplib)  Gmail SMTP reminders'
)

doc.add_heading('3.2 Request Flow for AI Chat', level=2)
steps = [
    'User types a message in the chat interface',
    'Browser sends POST /api/chat with the message',
    'app.py loads user profile from SQLite (name, goal, age, work/free time)',
    'RLHF StrategySelector picks a coaching style (epsilon-greedy: 80% exploit, 20% explore)',
    'Dynamic persona is built based on user\'s goal (fitness coach, dev mentor, etc.)',
    'System prompt + last 10 chat messages + current message  sent to llama-server',
    'llama-server streams response token-by-token via Server-Sent Events (SSE)',
    'Frontend renders markdown in real-time as tokens arrive',
    'Full response saved to chat_history table in SQLite',
    'Smart suggestion poll buttons generated based on response content',
    'User can click thumbs up/down  RLHF feedback logged and strategy scores updated',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: ')
    run.bold = True
    p.add_run(step)

doc.add_heading('3.3 Component Files', level=2)
add_table(
    ['File', 'Lines', 'Purpose'],
    [
        ['web/app.py', '~3300', 'Main Flask server  routes, API, chat logic, commands'],
        ['memory.py', '~1080', 'SQLite data layer  all DB operations, 18+ tables'],
        ['local_llm.py', '~106', 'LLM wrapper  connects to llama-server, streaming, retries'],
        ['reminders.py', '~80', 'Time parser  "10m", "5pm", "in 2 hours" etc.'],
        ['ai_task_scheduler.py', '~183', 'APScheduler  3 cron jobs for automated emails'],
        ['habit_intelligence.py', '~40', 'Habit analysis  completion stats and insights'],
        ['coach_engine.py', '~25', 'Weekly report generation  scores, wins, improvements'],
        ['smart_blocks.py', '', 'Knowledge block management  CRUD, linking, suggestions'],
        ['rlhf/strategy_selector.py', '~80', 'RLHF  epsilon-greedy strategy selection'],
        ['rlhf/reward_engine.py', '~25', 'RLHF  feedback score calculation'],
        ['rlhf/feedback_manager.py', '~35', 'RLHF  orchestrates feedback processing'],
        ['rlhf/storage.py', '~82', 'RLHF  SQLite tables for logs and scores'],
    ]
)

doc.add_page_break()

# 
# 4. AI FINE-TUNING  RLHF SYSTEM
# 
doc.add_heading('4. AI Fine-Tuning  RLHF System (Step by Step)', level=1)

doc.add_paragraph(
    'PartnerAI does NOT fine-tune the model weights (Phi-3 runs as-is from the GGUF file). '
    'Instead, it uses Reinforcement Learning from Human Feedback (RLHF) at the PROMPT level  '
    'adjusting HOW the AI talks (coaching style), not WHAT it knows. This approach requires '
    'no GPU for training and adapts instantly based on user feedback.'
)

doc.add_heading('Step 1: Define Coaching Strategies', level=2)
doc.add_paragraph('File: rlhf/strategy_selector.py')
doc.add_paragraph(
    'Six distinct coaching strategies are defined, each changing how the AI responds:'
)
add_table(
    ['Strategy', 'System Prompt Instruction'],
    [
        ['direct_action', 'Provide a direct, no-nonsense answer. Focus on immediate action.'],
        ['step_by_step', 'Break down the solution into clear, numbered steps.'],
        ['deep_explanation', 'Provide comprehensive explanation. Cover the "why" and "how" in depth.'],
        ['technical_breakdown', 'Focus on technical details, code structure, and mechanics.'],
        ['motivational_push', 'Be encouraging and high-energy. Use emojis. Motivate action.'],
        ['strategic_analysis', 'Analyze from a high-level perspective. Discuss pros, cons, long-term.'],
    ]
)

doc.add_heading('Step 2: Strategy Selection (Epsilon-Greedy Algorithm)', level=2)
doc.add_paragraph('File: rlhf/strategy_selector.py  get_best_strategy()')
doc.add_paragraph(
    'Before every AI response, the system selects which coaching strategy to use:'
)
add_bullet('Load all strategy scores from the rlhf_strategy_scores database table')
add_bullet('If last strategy received negative feedback  remove it from candidates')
add_bullet('80% of the time  Pick the highest-scoring strategy (EXPLOITATION)')
add_bullet('20% of the time  Pick a random strategy (EXPLORATION)')
add_bullet('This is the epsilon-greedy algorithm from reinforcement learning')

doc.add_paragraph(
    'This balances between using what works (exploit) and trying new approaches (explore), '
    'ensuring the system doesn\'t get stuck in a local optimum.'
)

doc.add_heading('Step 3: Inject Strategy into System Prompt', level=2)
doc.add_paragraph('File: web/app.py  chat handler (~line 1245)')
doc.add_paragraph(
    'The selected strategy becomes a coaching instruction inside the AI\'s system prompt:'
)
add_code_block(
    'system_prompt = f"""{persona}\n\n'
    'USER CONTEXT:\n'
    '- Name: {name}, Age: {age}\n'
    '- Main goal: {goal}\n'
    '- Coaching style this session: {strategy_instruction}\n'
    '..."""'
)
doc.add_paragraph(
    'For example, if "motivational_push" is selected, the AI receives: '
    '"Be encouraging and high-energy. Focus on motivating the user to take action. Use emojis."'
)

doc.add_heading('Step 4: User Gives Feedback', level=2)
doc.add_paragraph('File: web/templates/chat.html  HITL (Human-in-the-Loop) bar')
doc.add_paragraph(
    'After every AI response in the chat, the user sees action buttons:'
)
add_bullet(' Thumbs Up  marks the response as helpful')
add_bullet(' Thumbs Down  marks the response as not helpful')
add_bullet(' Correct  allows the user to submit a text correction')
doc.add_paragraph(
    'When clicked, a POST request is sent to /api/feedback with the user input, '
    'AI response, strategy used, and feedback label.'
)

doc.add_heading('Step 5: Score Calculation', level=2)
doc.add_paragraph('File: rlhf/reward_engine.py')
doc.add_paragraph('The feedback label is converted to a numeric score:')
add_table(
    ['Feedback Label', 'Numeric Score'],
    [
        ['very_helpful', '+2 points'],
        ['helpful', '+1 point'],
        ['not_helpful', '-1 point'],
    ]
)

doc.add_heading('Step 6: Feedback Processing & Storage', level=2)
doc.add_paragraph('File: rlhf/feedback_manager.py  process_feedback()')
doc.add_paragraph('When feedback is received, the FeedbackManager does 4 things:')

p = doc.add_paragraph()
run = p.add_run('1. Calculate Score: ')
run.bold = True
p.add_run('Convert label to number via RewardEngine')

p = doc.add_paragraph()
run = p.add_run('2. Log to Database: ')
run.bold = True
p.add_run('Insert into rlhf_feedback_logs table (user_input, ai_response, strategy, score, timestamp)')

p = doc.add_paragraph()
run = p.add_run('3. Update Strategy Score: ')
run.bold = True
p.add_run('UPDATE rlhf_strategy_scores SET total_score = total_score + score_delta')

p = doc.add_paragraph()
run = p.add_run('4. Handle Negative: ')
run.bold = True
p.add_run('If score < 0, tell StrategySelector to avoid this strategy next turn')

doc.add_heading('Step 7: The Feedback Loop Closes', level=2)
doc.add_paragraph(
    'Next time the user sends a message, Step 2 runs again  but now the strategy scores '
    'have been updated by the feedback. Over time, strategies that get more positive feedback '
    'accumulate higher scores and are selected more often. Bad strategies get penalized and '
    'naturally phase out.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Result: ')
run.bold = True
p.add_run(
    'Without any model retraining or GPU compute, PartnerAI adapts its coaching personality '
    'to each user\'s preferences. A user who prefers step-by-step explanations will gradually '
    'see more of that style. A user who responds well to motivation will get more energetic responses.'
)

doc.add_page_break()

# 
# 5. DYNAMIC PERSONA SYSTEM
# 
doc.add_heading('5. Dynamic Persona System', level=1)

doc.add_paragraph(
    'The AI doesn\'t have a static personality. It builds a domain-specific persona '
    'based on the user\'s career goal, stored during onboarding.'
)

add_table(
    ['User Goal Contains', 'AI Persona', 'Personality'],
    [
        ['fitness, gym, workout', 'Fitness Coach', 'Tough love + genuine care, discipline over motivation'],
        ['code, programming, python', 'Senior Dev Mentor', 'Debug thinking patterns, celebrate shipping'],
        ['business, startup, freelance', 'Business Mentor', 'Revenue first, cut through overthinking'],
        ['study, exam, college', 'Study Coach', 'Evidence-based techniques, energy management'],
        ['(anything else)', 'Personal Growth Mentor', 'Build habits and momentum, specific advice'],
    ]
)

doc.add_paragraph('Each persona also receives injected context:')
add_bullet('User\'s name, age, and career goal')
add_bullet('Work time and free time schedule')
add_bullet('Current RLHF strategy instruction')
add_bullet('Conversation rules (max 1 question per reply, no filler openers, etc.)')

doc.add_page_break()

# 
# 6. AUTOMATION SYSTEM
# 
doc.add_heading('6. Automation System', level=1)

doc.add_heading('6.1 Background Scheduler', level=2)
doc.add_paragraph('File: ai_task_scheduler.py')
doc.add_paragraph(
    'Uses APScheduler BackgroundScheduler with CronTrigger. '
    'Starts automatically when app.py runs. Three cron jobs:'
)

add_table(
    ['Job', 'Time (IST)', 'What It Does'],
    [
        ['Morning Reminders', '8:30 AM', 'Emails today\'s AI tasks + motivational quote to all active users'],
        ['Evening Reminders', '8:00 PM', 'Emails incomplete tasks + encouragement message'],
        ['Midnight Rollover', '12:00 AM', 'Moves unfinished tasks to tomorrow automatically'],
    ]
)

doc.add_heading('How it starts:', level=3)
add_code_block(
    '# In app.py __main__ block:\n'
    'from ai_task_scheduler import init_scheduler\n'
    'scheduler = init_scheduler(app)\n'
    'print(" AI Task Scheduler initialized")'
)

doc.add_heading('6.2 Email System', level=2)
add_table(
    ['Setting', 'Value'],
    [
        ['SMTP Server', 'smtp.gmail.com'],
        ['Port', '587 (TLS)'],
        ['From Address', 'dreamsyncai07@gmail.com'],
        ['Authentication', 'App-specific password'],
        ['Library', 'smtplib + email.mime (Python standard library)'],
    ]
)

doc.add_page_break()

# 
# 7. REMINDER & NOTIFICATION SYSTEM
# 
doc.add_heading('7. Reminder & Notification System', level=1)

doc.add_heading('7.1 Three Ways to Set Reminders', level=2)

p = doc.add_paragraph()
run = p.add_run('1. /reminder command: ')
run.bold = True
p.add_run('/reminder 10m Drink water  or  /reminder 5pm Call mom')

p = doc.add_paragraph()
run = p.add_run('2. Natural language: ')
run.bold = True
p.add_run('"remind me to exercise in 30 minutes"')

p = doc.add_paragraph()
run = p.add_run('3. Natural language: ')
run.bold = True
p.add_run('"remind me to call mom at 5pm"')

doc.add_heading('7.2 Time Parser (reminders.py)', level=2)
doc.add_paragraph('Supports all these formats:')
add_table(
    ['Input', 'Parsed As'],
    [
        ['10m', '10 minutes (timedelta)'],
        ['2h', '2 hours (timedelta)'],
        ['30s', '30 seconds (timedelta)'],
        ['5pm', 'Today 5:00 PM or tomorrow if past (datetime)'],
        ['5:30pm', 'Today 5:30 PM (datetime)'],
        ['at 5pm', 'Today 5:00 PM (datetime)'],
        ['10 min', '10 minutes (timedelta)'],
        ['in 10 minutes', '10 minutes (timedelta)'],
        ['2hr', '2 hours (timedelta)'],
        ['30sec', '30 seconds (timedelta)'],
    ]
)

doc.add_heading('7.3 What Happens When a Reminder Fires', level=2)
add_bullet('Email sent via Gmail SMTP to user\'s registered email address')
add_bullet('In-app reminder stored in SQLite reminders table')
add_bullet('Frontend polls /api/reminders/check every 15 seconds')
add_bullet('Toast popup slides in from top-right with bell animation')
add_bullet('Reminder message injected into chat as AI message')
add_bullet('User can dismiss notification (calls /api/reminders/dismiss)')

doc.add_page_break()

# 
# 8. DATABASE SCHEMA
# 
doc.add_heading('8. Database Schema', level=1)

doc.add_paragraph('File: memory.py  SQLite database with 18+ tables')

doc.add_heading('8.1 Core Tables', level=2)
add_table(
    ['Table', 'Purpose', 'Key Columns'],
    [
        ['users', 'User profiles & auth', 'user_id, name, career, email, password, state, streak'],
        ['chat_history', 'All AI conversations', 'user_id, role, content, timestamp'],
        ['ai_daily_tasks', 'AI-generated tasks', 'user_id, task_content, task_date, status'],
        ['daily_tasks', 'User-created tasks', 'user_id, task_content, is_completed'],
        ['habits', 'Habit definitions', 'user_id, title, frequency, time_of_day'],
        ['habit_logs', 'Completion logs', 'habit_id, log_date, status'],
        ['reminders', 'In-app notifications', 'user_id, content, trigger_at, dismissed'],
        ['smart_blocks', 'Knowledge notes', 'user_id, block_type, title, content, metadata'],
        ['block_relationships', 'Links between blocks', 'block_id_1, block_id_2, type'],
        ['focus_sessions', 'Pomodoro logs', 'user_id, duration, task, completed_at'],
        ['user_rewards', 'Gamification', 'user_id, reward_type, earned_at'],
        ['ai_user_memory', 'AI memory of user', 'user_id, memory_key, memory_value, confidence'],
    ]
)

doc.add_heading('8.2 Team Collaboration Tables', level=2)
add_table(
    ['Table', 'Purpose'],
    [
        ['groups', 'Team groups with name, goal, invite code, deadline'],
        ['group_members', 'Membership with roles (leader/member)'],
        ['group_tasks', 'Assigned tasks with status tracking'],
        ['group_chat_messages', 'Team chat history'],
    ]
)

doc.add_heading('8.3 RLHF Tables', level=2)
add_table(
    ['Table', 'Purpose', 'Key Columns'],
    [
        ['rlhf_feedback_logs', 'Every feedback event', 'user_input, ai_response, strategy_type, feedback_label, numeric_score'],
        ['rlhf_strategy_scores', 'Aggregated scores', 'strategy_type, total_score, usage_count'],
    ]
)

doc.add_heading('8.4 Users Table Schema (19 columns)', level=2)
add_table(
    ['Index', 'Column', 'Type', 'Description'],
    [
        ['0', 'user_id', 'INTEGER PK', 'Primary key'],
        ['1', 'name', 'TEXT', 'Display name'],
        ['2', 'career', 'TEXT', 'Goal / career objective'],
        ['3', 'hobbies', 'TEXT', 'User hobbies'],
        ['4', 'last_task', 'TEXT', 'Current focus task'],
        ['5', 'task_status', 'TEXT', 'pending/completed'],
        ['6', 'state', 'TEXT', 'ACTIVE/ONBOARDING'],
        ['7', 'tasks_completed', 'INTEGER', 'Total tasks done'],
        ['8', 'streak', 'INTEGER', 'Current streak days'],
        ['9', 'last_active_date', 'TEXT', 'Last login date'],
        ['10', 'daily_topic', 'TEXT', 'Today\'s focus topic'],
        ['11', 'work_time', 'TEXT', 'Work schedule'],
        ['12', 'free_time', 'TEXT', 'Free time available'],
        ['13', 'age', 'TEXT', 'User age'],
        ['14', 'last_task_date', 'TEXT', 'Last task generation date'],
        ['15', 'username', 'TEXT', 'Login username'],
        ['16', 'password', 'TEXT', 'Login password'],
        ['17', 'email', 'TEXT', 'Email for reminders'],
        ['18', 'flow_day', 'INTEGER', 'Flow content day counter'],
    ]
)

doc.add_page_break()

# 
# 9. FRONTEND PAGES
# 
doc.add_heading('9. Frontend Pages', level=1)

add_table(
    ['Page', 'Template', 'Features'],
    [
        ['Login', 'login.html', 'Username/password authentication'],
        ['Onboarding', 'onboarding.html', 'Multi-step wizard (name, goal, schedule)'],
        ['Home', 'home.html', 'Dashboard with stats, streak, quick actions'],
        ['AI Chat', 'chat.html', 'Streaming chat, poll suggestions, feedback, deep think'],
        ['Habits', 'habits.html', 'Habit tracker with AI insights'],
        ['Focus Mode', 'focus_mode.html', 'Pomodoro timer with music/wallpapers'],
        ['Knowledge Blocks', 'workspace.html', 'Notion-like card grid for notes'],
        ['Block Editor', 'block_editor.html', 'Rich editing for knowledge blocks'],
        ['Report', 'report.html', 'Weekly progress analytics'],
        ['Settings', 'settings.html', 'Profile, theme (dark/light), preferences'],
        ['Community', 'community.html', 'Social feed and posts'],
        ['Team', 'group.html', 'Team collaboration, shared tasks, team chat'],
    ]
)

doc.add_heading('Design System', level=2)
doc.add_paragraph('File: web/static/css/design-system.css')
add_table(
    ['Property', 'Dark Theme Value'],
    [
        ['Background', '#121212'],
        ['Surface', '#1E1E24'],
        ['Elevated', '#2C2C35'],
        ['Accent (matte purple)', '#8B78CC'],
        ['Accent hover', '#7A66C2'],
        ['Text primary', '#E0E0E0'],
        ['Text secondary', '#B0B0B0'],
    ]
)

doc.add_page_break()

# 
# 10. CHAT COMMAND SYSTEM
# 
doc.add_heading('10. Chat Command System', level=1)

add_table(
    ['Command', 'What It Does'],
    [
        ['/daily', 'AI generates 3-5 personalized tasks for today'],
        ['/plan', 'Creates a structured daily plan with time blocks'],
        ['/reminder 10m msg', 'Schedules email + in-app reminder'],
        ['/custom tasks', 'Sets custom focus tasks manually'],
        ['/report', 'Shows productivity statistics and progress'],
        ['/article', 'AI writes a motivational article for your goal'],
        ['/news', 'Fetches curated news based on user\'s career goal'],
        ['/question query', 'General knowledge Q&A'],
        ['/help decision', 'Gets decisive recommendation for a decision'],
        ['/mt', 'Motivational content and quotes'],
        ['/task topic', 'Sets daily learning topic'],
        ['/reset', 'Wipes all data and restarts onboarding'],
    ]
)

doc.add_heading('Poll Suggestion System', level=2)
doc.add_paragraph(
    'After every AI response, smart clickable suggestion buttons appear automatically. '
    'The system detects the context of the AI\'s message and generates relevant quick replies:'
)
add_table(
    ['AI Response Context', 'Suggestion Buttons Shown'],
    [
        ['Asks a yes/no question', 'Yes | No | Tell me more'],
        ['Asks "are you ready?"', 'Yes, let\'s go!  | Not yet | Tell me more'],
        ['Talks about tasks/plans', 'Show my tasks | What\'s next? | /daily'],
        ['Sets a reminder', 'Set another reminder | Show my tasks | Thanks! '],
        ['Gives motivation', 'What\'s next? | Give me a challenge | Thanks! '],
        ['Discusses code', 'Show me the code | Explain more | What else?'],
        ['Welcome/greeting', '/daily | /plan | I need help with something'],
        ['Default', 'Tell me more | What\'s next? | /daily'],
    ]
)

doc.add_page_break()

# 
# 11. HOW TO RUN
# 
doc.add_heading('11. How to Run the Project', level=1)

steps = [
    ('Start the LLM Server',
     'E:\\llama_cpp\\llama-server.exe -m E:\\llama_cpp\\Phi-3-mini-4k-instruct-q4.gguf -c 2048 --host 127.0.0.1 --port 8080'),
    ('Activate Python Environment',
     'cd E:\\PartnerAI\n.venv\\Scripts\\Activate.ps1'),
    ('Install Dependencies',
     'pip install -r requirements.txt'),
    ('Start Flask Server',
     'cd web\npython app.py'),
    ('Open in Browser',
     'http://127.0.0.1:5000'),
]

for i, (title, cmd) in enumerate(steps, 1):
    doc.add_heading(f'Step {i}: {title}', level=2)
    add_code_block(cmd)

doc.add_heading('What Happens on Startup', level=2)
add_bullet('All database tables initialized (memory.py init_db)')
add_bullet('APScheduler starts with 3 background cron jobs')
add_bullet('llama-server auto-started if not already running')
add_bullet('Flask serves the web app on port 5000 (all interfaces)')

doc.add_page_break()

# 
# 12. KEY DESIGN DECISIONS
# 
doc.add_heading('12. Key Design Decisions', level=1)

add_table(
    ['Decision', 'Reason'],
    [
        ['Local LLM (Phi-3) instead of OpenAI API', 'Privacy, no API costs, works offline, no rate limits'],
        ['llama.cpp instead of Ollama', 'More control, lighter footprint, direct GGUF loading'],
        ['RLHF at prompt level instead of model fine-tuning', 'No GPU needed, instant adaptation, simple implementation'],
        ['SQLite instead of PostgreSQL', 'Zero setup, single-file DB, perfect for single-user/small scale'],
        ['Streaming responses (SSE)', 'Better UX  user sees text appear word-by-word'],
        ['Threading for reminders', 'Non-blocking  server stays responsive while timer waits'],
        ['APScheduler in-process', 'No external cron daemon needed, runs inside Flask'],
        ['Vanilla JS instead of React', 'No build step, simpler deployment, lighter bundle'],
        ['Dark theme by default', 'Reduces eye strain for productivity app used daily'],
        ['Epsilon-greedy for RLHF', 'Balances exploitation of best strategy with exploration of alternatives'],
    ]
)

doc.add_page_break()

# 
# 13. FILE STRUCTURE
# 
doc.add_heading('13. File Structure Reference', level=1)

add_code_block(
    'E:\\PartnerAI\\\n'
    ' web/\n'
    '    app.py                  # Main Flask server (~3300 lines)\n'
    '    static/\n'
    '       css/design-system.css   # Theme & design tokens\n'
    '       js/focus_mode_v5.js     # Focus mode logic\n'
    '       images/                 # Wallpapers, icons\n'
    '    templates/\n'
    '        base.html               # Base layout + nav\n'
    '        chat.html               # AI chat interface\n'
    '        home.html               # Dashboard\n'
    '        habits.html             # Habit tracker\n'
    '        workspace.html          # Knowledge blocks\n'
    '        block_editor.html       # Block editor\n'
    '        focus_mode.html         # Pomodoro timer\n'
    '        report.html             # Weekly report\n'
    '        settings.html           # User settings\n'
    '        login.html              # Authentication\n'
    '        onboarding.html         # Setup wizard\n'
    '        community.html          # Social feed\n'
    '        group.html              # Team page\n'
    ' rlhf/\n'
    '    strategy_selector.py    # Epsilon-greedy selection\n'
    '    reward_engine.py        # Score calculation\n'
    '    feedback_manager.py     # Feedback orchestration\n'
    '    storage.py              # RLHF database tables\n'
    ' memory.py                   # SQLite data layer\n'
    ' local_llm.py                # LLM wrapper (llama-server)\n'
    ' reminders.py                # Time parser\n'
    ' ai_task_scheduler.py        # APScheduler cron jobs\n'
    ' habit_intelligence.py       # Habit analysis\n'
    ' coach_engine.py             # Weekly reports\n'
    ' smart_blocks.py             # Knowledge blocks logic\n'
    ' requirements.txt            # Python dependencies\n'
    ' partnerai.db                # SQLite database\n'
    ' start_llama.bat             # LLM server launcher'
)

# 
# SAVE
# 
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'PartnerAI_Project_Report.docx')
doc.save(output_path)
print(f"\n Report saved to: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024:.1f} KB")
